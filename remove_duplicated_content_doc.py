import sys
from docx import Document
import locale
import pypinyin
from pykakasi import kakasi
import re

def get_pinyin(text):
    # Convert Chinese text to pinyin
    pinyin_list = pypinyin.lazy_pinyin(text)
    return ''.join(pinyin_list)

def extract_romaji(text):
    # Extract romaji enclosed in parentheses
    romaji_pattern = r'\((.*?)\)'
    match = re.search(romaji_pattern, text)
    if match:
        return match.group(1).strip()
    else:
        return None

def get_romaji(text):
    # Convert Japanese text to romaji using pykakasi
    k = kakasi()
    k.setMode('H', 'a')  # Hiragana to ascii
    k.setMode('K', 'a')  # Katakana to ascii
    k.setMode('J', 'a')  # Japanese character to ascii
    conv = k.getConverter()
    return conv.do(text)

def remove_duplicates_and_empty_lines(input_file):
    # Open the Word document
    doc = Document(input_file)
    
    # List to store unique lines
    unique_lines = set()
    
    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        # Get the text of the paragraph
        text = para.text.strip()
        
        # Check if the line is not empty and not a duplicate
        if text:
            unique_lines.add(text)
    
    # Sort the unique lines
    sorted_lines = sorted(
        unique_lines,
        key=lambda x: extract_romaji(x) if extract_romaji(x) is not None
        else get_pinyin(x) if any('\u4e00' <= char <= '\u9fff' for char in x)
        else get_romaji(x) if any('\u3040' <= char <= '\u30ff' for char in x)
        else locale.strxfrm(x)
    )
    
    # Create a new document to save unique lines
    new_doc = Document()
    
    # Add sorted lines to the new document
    for line in sorted_lines:
        new_doc.add_paragraph(line)
    
    # Save the new document
    new_filename = input_file.split(".docx")[0] + "_output.docx"
    new_doc.save(new_filename)
    
    return new_filename

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python script.py input.docx")
    else:
        # Set locale to handle sorting of different languages correctly
        locale.setlocale(locale.LC_ALL, '')
        
        input_file = sys.argv[1]
        
        # Remove duplicates and empty lines, and save to new file
        output_file = remove_duplicates_and_empty_lines(input_file)
        print(f"Output saved to {output_file}")
